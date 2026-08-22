"""
Load Word (.docx) files into RawDocument objects using `python-docx`.
"""

from __future__ import annotations

import logging
import os

from ingestion.schema import RawDocument, stable_document_id

logger = logging.getLogger(__name__)


def load_docx(file_path: str) -> RawDocument:
    try:
        import docx  # python-docx
    except ImportError as exc:  # pragma: no cover
        raise ImportError(
            "python-docx is required for DOCX ingestion. Install it with `pip install python-docx`."
        ) from exc

    if not os.path.isfile(file_path):
        raise FileNotFoundError(f"DOCX not found: {file_path}")

    document = docx.Document(file_path)

    paragraphs = [p.text for p in document.paragraphs if p.text and p.text.strip()]

    # Also pull table content, since DOCX knowledge bases often use tables
    # for structured facts (pricing, specs, policy tables, etc).
    table_lines: list[str] = []
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                table_lines.append(" | ".join(cells))

    full_text = "\n".join(paragraphs)
    if table_lines:
        full_text += "\n\n" + "\n".join(table_lines)

    core_props = document.core_properties
    title = (core_props.title or os.path.basename(file_path)).strip()

    document_id = stable_document_id("docx", file_path)

    return RawDocument(
        document_id=document_id,
        source_type="docx",
        source_ref=file_path,
        title=title,
        content=full_text,
        metadata={
            "source_type": "docx",
            "file_name": os.path.basename(file_path),
            "author": core_props.author,
            "has_tables": bool(table_lines),
        },
    )


def load_docx_from_dir(directory: str) -> list[RawDocument]:
    if not os.path.isdir(directory):
        return []
    documents: list[RawDocument] = []
    for name in sorted(os.listdir(directory)):
        if name.lower().endswith(".docx"):
            path = os.path.join(directory, name)
            try:
                documents.append(load_docx(path))
            except Exception as exc:  # noqa: BLE001
                logger.error("Skipping unreadable DOCX %s: %s", path, exc)
    return documents
